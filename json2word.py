"""
从json中读取数据挖掘的作业，并且将其写入word中
"""
import json
import re
import requests
import docx
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from PIL import Image
import os
import base64


def get_json(path):
    """
    读取json文件，并将其返回
    :param path:
    :return:
    """
    with open(path, 'r') as f:
        result = json.load(f)
    return result


def deal_json(data, document):
    """
    处理json文件中的问题
    :param data: json数据
    :param document: 要写入的文档
    :return:
    """
    problems = data['data']['problems']
    head_name = data['data']['name']
    print(head_name)
    document.add_heading(head_name, level=0)
    for index, problem in enumerate(problems):
        type = problem['content']['Type']
        if type == 'SingleChoice':
            deal_SingleChoice(index, problem, document)
        elif type == 'MultipleChoice':
            deal_MultipleChoice(index, problem, document)
        elif type == 'FillBlank':
            deal_FillBlank(index, problem, document)
        elif type == 'Judgement':
            deal_Judgement(index, problem, document)
        else:
            print("该问题没有处理方法", type)


def deal_SingleChoice(index, problem, document):
    """
    单选
    :param index:   序号
    :param problem: 输入问题的json
    :param document: 写入的word文件
    :return:
    """
    retsult = ""
    question = problem['content']['Body']
    options = problem['content']['Options']
    type = problem['content']['TypeText']
    p = document.add_paragraph(str(index + 1) + ". 【" + type + "】")
    deal_string(question, p, document)
    for option in options:
        p = document.add_paragraph(option['key'] + ". ")
        deal_string(option['value'], p, document)

    if 'answer' in problem['user']:
        anwser = problem['user']['answer']
        document.add_paragraph("[答案] " + anwser[0])
    else:
        print("第%d道单选题没有答案：" % (index + 1))
    document.add_paragraph()


def deal_MultipleChoice(index, problem, document):
    """
    多选题
    :param index:   序号
    :param problem: 输入问题的json
    :param document: 写入的word文件
    :return:
    """
    question = problem['content']['Body']
    options = problem['content']['Options']

    type = problem['content']['TypeText']
    p = document.add_paragraph(str(index + 1) + ". 【" + type + "】")
    deal_string(question, p, document)
    for option in options:
        # document.add_paragraph(option['key'] + ". " + rm_tag(option['value']))
        p = document.add_paragraph(option['key'] + ". ")
        deal_string(option['value'], p, document)

    p = document.add_paragraph("[答案] ")
    if 'answer' in problem['user']:
        answers = problem['user']['answer']
        for answer in answers:
            p.add_run(answer + " ")
    else:
        print("第%d道多选题没有答案：" % (index + 1))
    document.add_paragraph()


def deal_FillBlank(index, problem, document):
    """
    填空题
    :param index:   序号
    :param problem: 输入问题的json
    :param document: 写入的word文件
    :return:
    """
    question = problem['content']['Body']
    type = problem['content']['TypeText']
    p = document.add_paragraph(str(index + 1) + ". 【" + type + "】")
    deal_string(question, p, document)
    p = document.add_paragraph("[答案] ")

    if 'answers' in problem['user']:
        answers = problem['user']['answers']
        for answer in answers:
            # for an in answers[answer]:
            #     p.add_run(an + ",")
            p.add_run(",".join(answers[answer]))
            p.add_run(";  ")
    else:
        print("第%d道填空题没有答案：" % (index + 1))
    document.add_paragraph()


def deal_Judgement(index, problem, document):
    """
    判断题
    :param index:   序号
    :param problem: 输入问题的json
    :param document: 写入的word文件
    :return:
    """
    question = problem['content']['Body']
    type = problem['content']['TypeText']
    p = document.add_paragraph(str(index + 1) + ". 【" + type + "】")
    deal_string(question, p, document)
    p = document.add_paragraph("[答案] ")
    if 'answer' in problem['user']:
        answer = problem['user']['answer'][0]
        if answer == "true":
            p.add_run("对")
        else:
            p.add_run("错")
    else:
        print("第%d道判断题没有答案：" % (index + 1))
    document.add_paragraph()


def deal_string(data, part, document):
    result = rm_tag(data)
    if test_img(result):
        get_img(result, part, document)
    else:
        part.add_run(result)


def test_img(data):
    """
    判断字符串中是否有图片
    :param data:
    :return:
    """
    img = r'<img.*?src="(.*?)".*?/>'
    if re.search(img, data):
        return True
    else:
        return False


def get_img(string, part, document):
    """
    读取图片，并将其保存
    :param string:
    :param part:
    :param document:
    :return:
    """
    img = r'<img.*?src="(.*?)".*?/>'
    urls = re.findall(img, string)
    # res = re.sub(img, "", string)
    r = part.add_run()
    sirs = string
    # r.add_run(res)
    res = re.search(r'<img.*?/>', sirs)
    while res is not None:
        if res.start() != 0:
            text = sirs[0:res.start()]
            sirs = sirs[res.start():]
            r.add_text(text)
        else:
            imgs = sirs[res.start():res.end()]
            sirs = sirs[res.end():]
            url = re.findall(img, imgs)[0]
            try:
                # 判断文件是否是base64编码
                if re.search(r'data:image/png', url):
                    url = re.sub(r'data:image/png;base64,', "", url)
                    imgdata = base64.b64decode(url)
                    fname = "base64.png"
                    with open("./img/" + fname, 'wb') as f:
                        f.write(imgdata)
                else:
                    response = requests.get(url)
                    fname = url.split("/")[-1]
                    if not re.search(r'.png$', fname):
                        fname += '.png'

                    if not os.path.exists("./img/"):
                        os.mkdir("./img/")

                    with open("./img/" + fname, 'wb') as f:
                        f.write(response.content)
                    
                im = Image.open("./img/" + fname)
                width = im.width
                # print("width="+str(width))
                if width > 500:
                    r.add_picture("./img/" + fname, width=Inches(6))
                else:
                    r.add_picture("./img/" + fname)

            except:
                print("读取的错误str：" + string)
                print(url + "文件读取失败！")
                pass

        res = re.search(r'<img.*?/>', sirs)

    if len(sirs) != 0:
        r.add_text(sirs)

    # for url in urls:
    #     try:
    #         response = requests.get(url)
    #         fname = url.split("/")[-1]
    #         if not re.search(r'.png$', fname):
    #             fname += '.png'
    #         with open("./img/" + fname, 'wb') as f:
    #             f.write(response.content)
    #         r.add_picture("./img/" + fname)
    #     except:
    #         print("读取的错误str：" + string)
    #         print(url + "文件读取失败！")
    #         pass


def rm_tag(data):
    """
    移除文件中的html标签
    :param data:
    :return:
    """
    p_tag = r'</?p.*?>'
    blank = r'&nbsp;'
    underline = r'<span.*?</span>'
    img = r'<img.*?src="(.*?)".*?/>'
    br = r'<br.*?/>'

    rdata = ""
    rdata = re.sub(p_tag, "", data)
    rdata = re.sub(blank, " ", rdata)
    rdata = re.sub(br, " ", rdata)
    rdata = re.sub(underline, "_____", rdata)
    rdata = re.sub(r'&gt;', ">", rdata)
    rdata = re.sub(r'&lt;', "<", rdata)
    rdata = re.sub(r'&quot;', "\"", rdata)
    # rdata = re.sub(img, "（这是一个图片）", rdata)

    return rdata


if __name__ == '__main__':
    document = Document()
    style = document.styles['Normal']
    font = style.font
    font.name = u'宋体'  # 设置为宋体
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

    # 读取多个文件进行尝试
    for i in range(1, 9):
        path = "./json/chapter0" + str(i) + ".json"
        result = get_json(path)
        deal_json(result, document)
    document.save("数据挖掘作业.docx")
